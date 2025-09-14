import os
import re
import json
from typing import List, Dict, Optional, Callable
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from .formatting_manager import FormattingManager


class Chapter:
    def __init__(self, title: str, content: str, chapter_number: int):
        self.title = title
        self.content = content
        self.chapter_number = chapter_number
        self.suggestions = []  # İşlenmiş öneriler (eski format)
        self.is_processed = False
        
        # Yeni alanlar - öneri geçmişi ve değişiklik takibi
        self.suggestion_history = []  # İşlenmiş öneriler
        self.last_modified = None     # Son değişiklik zamanı
        self.content_changes = []     # İçerik değişiklik geçmişi
        self.highlighting_info = {}   # Vurgulama bilgileri
        
        # YENİ: Beklemede olan öneriler için
        self.pending_suggestions = []  # Henüz işlem görmemiş öneriler
        
        # SIRALI ANALİZ DURUMU TAKİBİ - Her bölüm için hangi fazın tamamlandığını takip eder
        self.analysis_phases = {
            "grammar_completed": False,    # Dil Bilgisi fazı tamamlandı mı?
            "style_completed": False,      # Üslup fazı tamamlandı mı?
            "content_completed": False,    # İçerik fazı tamamlandı mı?
            "current_phase": "none"        # Mevcut faz: none, grammar, style, content, completed
        }
    
    def to_dict(self):
        # Suggestions listesini dict olarak serialize et
        suggestions_dict = []
        for suggestion in self.suggestions:
            if hasattr(suggestion, 'to_dict'):
                suggestions_dict.append(suggestion.to_dict())
            elif isinstance(suggestion, dict):
                suggestions_dict.append(suggestion)
            else:
                # Fallback: basit dict'e çevir
                suggestions_dict.append(str(suggestion))
        
        # Pending suggestions için de aynı işlemi yap
        pending_suggestions_dict = []
        for suggestion in getattr(self, 'pending_suggestions', []):
            if hasattr(suggestion, 'to_dict'):
                pending_suggestions_dict.append(suggestion.to_dict())
            elif isinstance(suggestion, dict):
                pending_suggestions_dict.append(suggestion)
            else:
                pending_suggestions_dict.append(str(suggestion))
        
        return {
            'title': self.title,
            'content': self.content,
            'chapter_number': self.chapter_number,
            'suggestions': suggestions_dict,
            'is_processed': self.is_processed,
            # Yeni alanlar
            'suggestion_history': getattr(self, 'suggestion_history', []),
            'last_modified': getattr(self, 'last_modified', None),
            'content_changes': getattr(self, 'content_changes', []),
            'highlighting_info': getattr(self, 'highlighting_info', {}),
            # YENİ: Beklemede olan öneriler
            'pending_suggestions': pending_suggestions_dict,
            # SIRALI ANALİZ DURUMU
            'analysis_phases': getattr(self, 'analysis_phases', {
                "grammar_completed": False,
                "style_completed": False,
                "content_completed": False,
                "current_phase": "none"
            })
        }
    
    @classmethod
    def from_dict(cls, data):
        chapter = cls(data['title'], data['content'], data['chapter_number'])
        
        # Suggestions listesini yükle - dict olarak kaydedilmiş olabilir
        suggestions_data = data.get('suggestions', [])
        chapter.suggestions = suggestions_data  # Dict olarak saklayacağız
        
        chapter.is_processed = data.get('is_processed', False)
        
        # Yeni alanları yükle
        chapter.suggestion_history = data.get('suggestion_history', [])
        chapter.last_modified = data.get('last_modified', None)
        chapter.content_changes = data.get('content_changes', [])
        chapter.highlighting_info = data.get('highlighting_info', {})
        
        # YENİ: Beklemede olan önerileri yükle
        chapter.pending_suggestions = data.get('pending_suggestions', [])
        
        # SIRALI ANALİZ DURUMU yükle
        chapter.analysis_phases = data.get('analysis_phases', {
            "grammar_completed": False,
            "style_completed": False,
            "content_completed": False,
            "current_phase": "none"
        })
        
        return chapter

class FileManager:
    def __init__(self):
        self.novel_path = None
        self.chapters = []
        self.novel_title = ""
        self.original_content = ""
        self.formatting_manager = FormattingManager()
    
    def load_novel(self, file_path: str, callback: Optional[Callable] = None):
        """Roman dosyasını yükle"""
        try:
            # Dosya uzantısına göre işlem yap
            if file_path.lower().endswith('.docx'):
                # Word dosyası yükleme
                content = self._load_docx_file(file_path)
            else:
                # TXT dosyası yükleme (mevcut davranış)
                with open(file_path, 'r', encoding='utf-8') as file:
                    content = file.read()
            
            self.novel_path = file_path
            self.novel_title = os.path.splitext(os.path.basename(file_path))[0]
            self.original_content = content
            
            if callback:
                callback(content)
            
            return True
        except Exception as e:
            print(f"Dosya yükleme hatası: {e}")
            return False
    
    def _load_docx_file(self, file_path: str) -> str:
        """Word dosyasını yükle ve biçimlendirmeleri koru"""
        try:
            doc = Document(file_path)
            content_parts = []
            
            for paragraph in doc.paragraphs:
                # Paragraf metnini, formatlamayı koruyarak 'run'lardan oluştur
                runs_text = []
                for run in paragraph.runs:
                    run_text = run.text
                    if run.bold:
                        run_text = f"{self.formatting_manager.inline_markers['bold']}{run_text}{self.formatting_manager.inline_markers['bold']}"
                    if run.italic:
                        run_text = f"{self.formatting_manager.inline_markers['italic']}{run_text}{self.formatting_manager.inline_markers['italic']}"
                    if run.underline:
                        run_text = f"{self.formatting_manager.inline_markers['underline']}{run_text}{self.formatting_manager.inline_markers['underline']}"
                    runs_text.append(run_text)
                
                paragraph_text = "".join(runs_text)
                
                # Paragraf seviyesi stiller (başlıklar)
                if paragraph.style and paragraph.style.name.lower().startswith(('heading', 'başlık')):
                    start_marker, end_marker = self.formatting_manager.paragraph_markers['heading']
                    paragraph_text = f"{start_marker}{paragraph_text}{end_marker}"
                
                # Paragraf hizalaması
                alignment = getattr(paragraph.paragraph_format.alignment, 'value', None)

                if alignment == 1:  # Orta
                    start_marker, end_marker = self.formatting_manager.paragraph_markers['centered']
                    paragraph_text = f"{start_marker}{paragraph_text}{end_marker}"
                elif alignment == 2:  # Sağ
                    start_marker, end_marker = self.formatting_manager.paragraph_markers['right_aligned']
                    paragraph_text = f"{start_marker}{paragraph_text}{end_marker}"
                # Not: Tkinter Text widget'ı 'iki yana yasla' (justify) özelliğini tam desteklemez,
                # bu yüzden şimdilik bunu sol hizalı olarak bırakıyoruz.
                
                content_parts.append(paragraph_text)
            
            return '\n'.join(content_parts)
        except Exception as e:
            print(f"Word dosyası yükleme hatası: {e}")
            # Hata durumunda basit metin olarak yükle
            doc = Document(file_path)
            return '\n'.join([paragraph.text for paragraph in doc.paragraphs])
    
    def split_into_chapters(self, content: str, method: str, custom_word: Optional[str] = None) -> List[Chapter]:
        """İçeriği bölümlere ayır"""
        lines = content.split('\n')
        chapters = []
        current_chapter_lines = []
        chapter_number = 1
        
        if method == "number_only":
            pattern = r'^\s*\d+\s*$'
        elif method == "keywords":
            # Yaygın bölüm başlık kelimeleri
            keywords = ["bölüm", "chapter", "kısım", "part", "fasıl"]
            pattern = r'^\s*(' + '|'.join(keywords) + r')\s*\d*\s*$'
        elif method == "custom" and custom_word:
            pattern = r'^\s*' + re.escape(custom_word) + r'\s*\d*\s*$'
        else:
            # Varsayılan: boş satırlarla ayır
            pattern = r'^\s*$'
        
        for line in lines:
            # Biçimlendirme etiketlerini temizle
            clean_line = self._remove_formatting_tags(line)
            
            if re.match(pattern, clean_line.strip(), re.IGNORECASE):
                # Yeni bölüm başlangıcı
                if current_chapter_lines:
                    chapter_content = '\n'.join(current_chapter_lines).strip()
                    if chapter_content:
                        chapter = Chapter(
                            title=f"Bölüm {chapter_number}",
                            content=chapter_content,
                            chapter_number=chapter_number
                        )
                        chapters.append(chapter)
                        chapter_number += 1
                    current_chapter_lines = []
                
                if method != "custom" or clean_line.strip():  # Boş satır değilse başlık olarak ekle
                    current_chapter_lines.append(line)
            else:
                current_chapter_lines.append(line)
        
        # Son bölümü ekle
        if current_chapter_lines:
            chapter_content = '\n'.join(current_chapter_lines).strip()
            if chapter_content:
                chapter = Chapter(
                    title=f"Bölüm {chapter_number}",
                    content=chapter_content,
                    chapter_number=chapter_number
                )
                chapters.append(chapter)
        
        self.chapters = chapters
        return chapters
    
    def _remove_formatting_tags(self, text: str) -> str:
        """Biçimlendirme etiketlerini temizle"""
        return self.formatting_manager.all_markers_regex.sub('', text)
    
    def get_chapter(self, chapter_number: int) -> Optional[Chapter]:
        """Belirli bir bölümü getir"""
        for chapter in self.chapters:
            if chapter.chapter_number == chapter_number:
                return chapter
        return None
    
    def update_chapter_content(self, chapter_number: int, new_content: str):
        """Bölüm içeriğini güncelle"""
        chapter = self.get_chapter(chapter_number)
        if chapter:
            chapter.content = new_content
    
    def _add_formatted_paragraph_to_docx(self, doc, text_line: str):
        """Parses a line of text with custom markers and adds it to the docx Document with proper formatting."""
        # 1. Handle paragraph-level formatting
        text_to_process = text_line
        alignment = WD_ALIGN_PARAGRAPH.LEFT
        style = None

        for key, (start_marker, end_marker) in self.formatting_manager.paragraph_markers.items():
            if text_line.startswith(start_marker) and text_line.endswith(end_marker):
                text_to_process = text_line[len(start_marker):-len(end_marker)]
                if key == 'heading':
                    style = 'Heading 2'
                elif key == 'centered':
                    alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif key == 'right_aligned':
                    alignment = WD_ALIGN_PARAGRAPH.RIGHT
                break

        p = doc.add_paragraph(style=style)
        p.alignment = alignment

        # 2. Handle inline formatting
        pattern = self.formatting_manager.all_markers_regex
        
        active_formats = set()
        last_pos = 0
        
        for match in pattern.finditer(text_to_process):
            start = match.start()
            if start > last_pos:
                segment = text_to_process[last_pos:start]
                if segment:
                    run = p.add_run(segment)
                    if 'bold' in active_formats:
                        run.bold = True
                    if 'italic' in active_formats:
                        run.italic = True
                    if 'underline' in active_formats:
                        run.underline = True
            
            marker = match.group(1)
            format_type = next((t for t, m in self.formatting_manager.inline_markers.items() if m == marker), None)
            if format_type and format_type in active_formats:
                active_formats.remove(format_type)
            else:
                active_formats.add(format_type)
            
            last_pos = match.end()
            
        if last_pos < len(text_to_process):
            segment = text_to_process[last_pos:]
            if segment:
                run = p.add_run(segment)
                if 'bold' in active_formats:
                    run.bold = True
                if 'italic' in active_formats:
                    run.italic = True
                if 'underline' in active_formats:
                    run.underline = True

    def export_novel(self, output_path: Optional[str] = None) -> Optional[str]:
        """Düzenlenmiş romanı dışa aktar"""
        if not output_path:
            output_path = f"{self.novel_title}_edited.txt"
        
        is_word_file = output_path.lower().endswith('.docx')
        
        if is_word_file:
            try:
                doc = Document()
                
                for chapter in sorted(self.chapters, key=lambda x: x.chapter_number):
                    paragraphs = chapter.content.split('\n')
                    for p_text in paragraphs:
                        if p_text.strip():
                            self._add_formatted_paragraph_to_docx(doc, p_text)
                        else:
                            doc.add_paragraph() # Preserve blank lines
                    
                    # Add a page break after each chapter
                    if chapter.chapter_number < len(self.chapters):
                        doc.add_page_break()

                doc.save(output_path)
                return output_path
            except Exception as e:
                print(f"Word dışa aktarma hatası: {e}")
                import traceback
                print(traceback.format_exc())
                return None
        else:
            # TXT dosyası olarak dışa aktar
            combined_content = []
            for chapter in sorted(self.chapters, key=lambda x: x.chapter_number):
                combined_content.append(chapter.content)
                combined_content.append("\n\n")
            
            final_content = '\n'.join(combined_content)
            
            try:
                with open(output_path, 'w', encoding='utf-8') as file:
                    file.write(final_content)
                return output_path
            except Exception as e:
                print(f"Dışa aktarma hatası: {e}")
                return None
    
    def save_chapters_to_json(self, file_path: str):
        """Bölümleri JSON formatında kaydet"""
        data = {
            'novel_title': self.novel_title,
            'novel_path': self.novel_path,
            'chapters': [chapter.to_dict() for chapter in self.chapters]
        }
        
        try:
            with open(file_path, 'w', encoding='utf-8') as file:
                json.dump(data, file, ensure_ascii=False, indent=2)
            return True
        except Exception as e:
            print(f"JSON kaydetme hatası: {e}")
            return False
    
    def load_chapters_from_json(self, file_path: str):
        """JSON'dan bölümleri yükle"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                data = json.load(file)
            
            self.novel_title = data.get('novel_title', '')
            self.novel_path = data.get('novel_path', '')
            self.chapters = [Chapter.from_dict(ch_data) for ch_data in data.get('chapters', [])]
            
            return True
        except Exception as e:
            print(f"JSON yükleme hatası: {e}")
            return False
    
    def get_state(self) -> Dict:
        """Mevcut durumu döndür"""
        return {
            'novel_title': self.novel_title,
            'novel_path': self.novel_path,
            'chapters': [chapter.to_dict() for chapter in self.chapters]
        }
    
    def load_state(self, state: Dict):
        """Durumu yükle"""
        self.novel_title = state.get('novel_title', '')
        self.novel_path = state.get('novel_path', '')
        self.chapters = [Chapter.from_dict(ch_data) for ch_data in state.get('chapters', [])]
        return self
